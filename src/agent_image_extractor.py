"""
Agent: Image Extractor
Extracts images and pages from PDF/DOCX documents for visual question reference.
Converts PDF pages to images (PNG) for display in assessment.
Extracts embedded images from DOCX files.
"""

import os
from pathlib import Path
from typing import List, Dict, Optional, Union


def extract_pdf_pages_as_images(
    pdf_path: str,
    output_dir: str,
    dpi: int = 150
) -> Dict[int, str]:
    """
    Extract PDF pages as images (PNG format) using PyMuPDF (fitz).
    
    Args:
        pdf_path: Path to PDF file
        output_dir: Directory to save images
        dpi: Resolution (150 = good balance of quality/size)
        
    Returns:
        Dict mapping page number -> image file path
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("ERROR: PyMuPDF not installed.")
        print("Install with: pip install PyMuPDF")
        return {}
    
    # Create output directory
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    try:
        # Open PDF
        print(f"Converting PDF to images: {pdf_path}")
        doc = fitz.open(pdf_path)
        
        page_images = {}
        zoom = dpi / 72  # Convert DPI to zoom factor (72 is default)
        
        for page_num in range(len(doc)):
            # Get page and render to image
            page = doc[page_num]
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # Save as PNG
            image_path = os.path.join(output_dir, f"page_{page_num + 1}.png")
            pix.save(image_path)
            page_images[page_num + 1] = image_path
            print(f"  [OK] Saved page {page_num + 1} -> {image_path}")
        
        doc.close()
        print(f"[OK] Extracted {len(page_images)} pages from PDF\n")
        return page_images
        
    except Exception as e:
        print(f"ERROR extracting PDF pages: {e}")
        return {}


def extract_pdf_images_only(
    pdf_path: str,
    output_dir: str
) -> List[str]:
    """
    Extract only embedded images from PDF (not page screenshots).
    More efficient but may miss images in some PDFs.
    
    Args:
        pdf_path: Path to PDF file
        output_dir: Directory to save images
        
    Returns:
        List of extracted image file paths
    """
    try:
        import PyPDF2
    except ImportError:
        print("ERROR: PyPDF2 not installed.")
        print("Install with: pip install PyPDF2")
        return []
    
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    try:
        print(f"Extracting embedded images from: {pdf_path}")
        image_paths = []
        
        with open(pdf_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                
                if "/XObject" in page["/Resources"]:
                    xObject = page["/Resources"]["/XObject"].get_object()
                    
                    image_count = 0
                    for obj in xObject:
                        if xObject[obj]["/Subtype"] == "/Image":
                            try:
                                size = (xObject[obj]["/Width"], xObject[obj]["/Height"])
                                data = xObject[obj].get_data()
                                
                                # Save as image file
                                image_path = os.path.join(
                                    output_dir,
                                    f"page_{page_num + 1}_image_{image_count}.png"
                                )
                                
                                # Create image from raw data
                                from PIL import Image
                                image = Image.new("RGB", size)
                                image.frombytes("RGB", size, data)
                                image.save(image_path)
                                
                                image_paths.append(image_path)
                                image_count += 1
                                print(f"  [OK] Extracted image from page {page_num + 1}")
                            except Exception as e:
                                print(f"  [WARNING] Could not extract image: {e}")
        
        print(f"[OK] Extracted {len(image_paths)} images\n")
        return image_paths
        
    except Exception as e:
        print(f"ERROR: {e}")
        return []


def extract_docx_images(
    docx_path: str,
    output_dir: str
) -> Dict[int, List[str]]:
    """
    Extract embedded images from DOCX file (python-docx).
    
    Args:
        docx_path: Path to DOCX file
        output_dir: Directory to save images
        
    Returns:
        Dict mapping paragraph index -> list of image file paths
    """
    try:
        from docx import Document
        from docx.oxml import parse_xml
        from io import BytesIO
    except ImportError:
        print("WARNING: python-docx not installed. Image extraction from DOCX skipped.")
        print("Install with: pip install python-docx")
        return {}
    
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    
    try:
        print(f"Extracting embedded images from DOCX: {docx_path}")
        doc = Document(docx_path)
        
        image_paths = {}
        image_count = 0
        
        # Extract images from document relationships
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image = rel.target_part.blob
                    
                    # Determine image format from content type
                    content_type = rel.target_part.content_type
                    ext = "png"
                    if "jpeg" in content_type or "jpg" in content_type:
                        ext = "jpg"
                    elif "png" in content_type:
                        ext = "png"
                    elif "gif" in content_type:
                        ext = "gif"
                    
                    # Save image
                    image_count += 1
                    image_path = os.path.join(output_dir, f"docx_image_{image_count}.{ext}")
                    
                    with open(image_path, "wb") as f:
                        f.write(image)
                    
                    if image_count not in image_paths:
                        image_paths[image_count] = []
                    image_paths[image_count].append(image_path)
                    
                    print(f"  [OK] Extracted image {image_count} -> {image_path}")
                except Exception as e:
                    print(f"  [WARNING] Could not extract embedded image: {e}")
        
        if image_count > 0:
            print(f"[OK] Extracted {image_count} images from DOCX\n")
        else:
            print(f"[INFO] No embedded images found in DOCX\n")
        
        return image_paths
        
    except Exception as e:
        print(f"ERROR extracting DOCX images: {e}")
        return {}


def docx_to_pdf_images(
    docx_path: str,
    output_dir: str,
    dpi: int = 150
) -> Dict[int, str]:
    """
    Convert DOCX to PDF first, then extract pages as images.
    Requires LibreOffice installed (for conversion).
    Fallback: returns empty dict if conversion fails.
    
    Args:
        docx_path: Path to DOCX file
        output_dir: Directory to save images
        dpi: Resolution for conversion
        
    Returns:
        Dict mapping page number -> image file path (or empty if conversion fails)
    """
    try:
        import subprocess
        import tempfile
    except ImportError:
        return {}
    
    try:
        # Try to convert DOCX to PDF using LibreOffice
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_pdf = os.path.join(temp_dir, "temp.pdf")
            
            # Try LibreOffice command
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf", "--outdir", temp_dir, docx_path],
                capture_output=True,
                timeout=30
            )
            
            if result.returncode == 0 and os.path.exists(temp_pdf):
                print(f"[INFO] Converted DOCX to PDF, extracting pages...")
                # Now extract pages from the converted PDF
                page_images = extract_pdf_pages_as_images(temp_pdf, output_dir, dpi)
                return page_images
            else:
                print(f"[INFO] LibreOffice conversion not available or failed")
                return {}
                
    except FileNotFoundError:
        print(f"[INFO] LibreOffice not found - skipping page image extraction for DOCX")
        return {}
    except Exception as e:
        print(f"[WARNING] Could not convert DOCX to PDF for images: {e}")
        return {}


def get_page_for_question(question_text: str, document_path: str) -> Optional[int]:
    """
    Determine which page a question comes from.
    Works with both PDF and DOCX files.
    Falls back to page 1 if not found.
    
    Args:
        question_text: The question text to find
        document_path: Path to PDF or DOCX file
        
    Returns:
        Page number (1-indexed), defaults to 1 if not found
    """
    file_ext = Path(document_path).suffix.lower()
    
    if file_ext == ".pdf":
        return _get_page_for_question_pdf(question_text, document_path)
    elif file_ext == ".docx":
        return _get_page_for_question_docx(question_text, document_path)
    else:
        return 1  # Default for unsupported formats


def _get_page_for_question_pdf(question_text: str, pdf_path: str) -> Optional[int]:
    """Find question page in PDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return 1
    
    try:
        doc = fitz.open(pdf_path)
        search_text = question_text[:100].lower()
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text().lower()
            
            if search_text[:30] in page_text:
                doc.close()
                return page_num + 1
        
        doc.close()
        return 1
        
    except Exception as e:
        print(f"[WARNING] Could not find question page in PDF: {e}")
        return 1


def _get_page_for_question_docx(question_text: str, docx_path: str) -> Optional[int]:
    """Find question location in DOCX (always returns 1 since DOCX is not paged)."""
    try:
        from docx import Document
    except ImportError:
        return 1
    
    try:
        doc = Document(docx_path)
        search_text = question_text[:100].lower()
        
        # Search through paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            if search_text[:30] in para.text.lower():
                return 1  # DOCX is continuous, return 1
        
        return 1
        
    except Exception as e:
        print(f"[WARNING] Could not find question in DOCX: {e}")
        return 1


if __name__ == "__main__":
    # Test extraction
    test_pdf = "data/input/PSLE_Math/PSLE-standard-math-paper-1-2-questions-with-answers-2024.pdf"
    
    if os.path.exists(test_pdf):
        output_dir = f"data/PSLE_Math_images"
        page_images = extract_pdf_pages_as_images(test_pdf, output_dir)
        print(f"Extracted {len(page_images)} pages")
    else:
        print(f"Test PDF not found: {test_pdf}")
